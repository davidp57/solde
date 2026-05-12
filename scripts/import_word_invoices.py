#!/usr/bin/env python3
"""Import historical client invoices from Word (.docx) files into Solde.

For each .docx file (pattern: "facture YYYY-NNNN.docx"):
  - Extract the invoice number from the filename.
  - Parse date, client name, address, and line items from the document.
  - Skip the file if an invoice with that number already exists in the DB.
  - Find or create the client contact.
  - Create the invoice with status=ARCHIVED (no accounting entries).
    - Convert the .docx to PDF and store it in data/pdfs/.
    - Attach the converted PDF through invoice.pdf_path.

Usage:
    python scripts/import_word_invoices.py \\
        --source /path/to/docx/folder [--db data/solde.db] [--commit] [--verbose]

Default mode is dry-run (no DB writes). Pass --commit to write to the database.
"""

from __future__ import annotations

import argparse
import re
import shutil
import sqlite3
import subprocess
import sys
from dataclasses import dataclass, field
from datetime import date
from decimal import Decimal, InvalidOperation
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any


# ---------------------------------------------------------------------------
# Patterns
# ---------------------------------------------------------------------------

# Matches "facture YYYY-NNNN.docx" → captures "YYYY-NNNN"
_FILENAME_RE = re.compile(r"^facture\s+(\d{4}-\d{4})\.docx$", re.IGNORECASE)

# Matches French date patterns: "01/04/2025", "1 avril 2025", "1er avril 2025", "2025-04-01"
_DATE_RE = re.compile(
    r"\b(\d{1,2})[/\-](\d{1,2})[/\-](\d{4})\b"
    r"|\b(\d{4})[/\-](\d{1,2})[/\-](\d{1,2})\b"
    r"|(?<!\d)(\d{1,2})\s*(?:er|[e\u00e8]me?|[e\u00e8]re?|[e\u00e8])?\s+(janvier|f[e\u00e9]vrier|mars|avril|mai|juin|juillet|ao[u\u00fb]t|septembre|octobre|novembre|d[e\u00e9]cembre)\s+(\d{4})\b",
    re.IGNORECASE,
)

_FRENCH_MONTHS = {
    "janvier": 1,
    "fevrier": 2,
    "février": 2,
    "mars": 3,
    "avril": 4,
    "mai": 5,
    "juin": 6,
    "juillet": 7,
    "aout": 8,
    "août": 8,
    "septembre": 9,
    "octobre": 10,
    "novembre": 11,
    "decembre": 12,
    "décembre": 12,
}

# Header/footer paragraphs to skip when looking for client name
_SKIP_PARAGRAPHS_RE = re.compile(
    r"^(les[\s\xa0]+etudes|metz,?\s+le\s|facture\s*n|en\s+votre|par\s+virement|bic\s*:|par\s+ch[eè]que|en\s+esp[eè]ces|r[eè]glement|pr[eé]cisez|pay[eé]|iban|rib\b)",
    re.IGNORECASE,
)

# Lines that look like a postal address (street), not a person's name
_STREET_START_RE = re.compile(
    r"^(\d+\s|résidence\b|rue\b|avenue\b|all[eé]e\b|impasse\b|boulevard\b|bd\b|chemins?\b|place\b|lotissements?\b|domaine\b|cité\b|voie\b|squares?\b|sq\b)",
    re.IGNORECASE,
)

# Salutation lines to skip (not a person's name)
_SALUTATION_RE = re.compile(
    r"^(m\.?\s*$|mme\.?\s*$|mr\.?\s*$|monsieur[\s,.:;-]*$|madame[\s,.:;-]*$|a\s+l.attention|objet\s*:|sujet\s*:|r[eé]f[eé]rence\s*:)",
    re.IGNORECASE,
)

# Price pattern: "26,00€" or "26.00" or "130 €"
_PRICE_RE = re.compile(r"([\d\s]+[,.][\d]+)\s*€?")

# Total row marker
_TOTAL_RE = re.compile(r"^total\s*$", re.IGNORECASE)


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class InvoiceLine:
    description: str
    amount: Decimal


@dataclass
class ParsedInvoice:
    number: str
    file: Path
    invoice_date: date | None
    client_name: str | None
    client_address: str | None
    lines: list[InvoiceLine]
    total_amount: Decimal


@dataclass
class ImportReport:
    created: list[str] = field(default_factory=list)
    skipped_exists: list[str] = field(default_factory=list)
    skipped_no_client: list[str] = field(default_factory=list)
    skipped_already_processed: list[str] = field(default_factory=list)
    pdf_attached: list[str] = field(default_factory=list)
    errors: list[tuple[str, str]] = field(default_factory=list)
    invoice_records: list[InvoiceRecord] = field(default_factory=list)


@dataclass
class InvoiceRecord:
    number: str
    client_name: str
    contact_created: bool
    year: int
    amount: Decimal
    invoice_date: date | None
    filename: str
    nb_lines: int


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _require_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)


def _normalize_name(text: str) -> str:
    """Lower-case, strip accents for fuzzy matching."""
    result = text.lower()
    for src, tgt in {
        "é": "e", "è": "e", "ê": "e", "ë": "e",
        "à": "a", "â": "a", "ä": "a",
        "î": "i", "ï": "i",
        "ô": "o", "ö": "o",
        "ù": "u", "û": "u", "ü": "u",
        "ç": "c",
    }.items():
        result = result.replace(src, tgt)
    return result


def _parse_price(text: str) -> Decimal | None:
    """Parse a price string like '26,00€' or '130 €' into a Decimal."""
    clean = text.replace("\xa0", "").replace(" ", "").replace("€", "").strip()
    clean = clean.replace(",", ".")
    try:
        return Decimal(clean).quantize(Decimal("0.01"))
    except InvalidOperation:
        return None


def _parse_date(text: str) -> date | None:
    """Extract a date from a text string."""
    m = _DATE_RE.search(text)
    if not m:
        return None
    groups = m.groups()
    try:
        if groups[0]:  # dd/mm/yyyy
            return date(int(groups[2]), int(groups[1]), int(groups[0]))
        if groups[3]:  # yyyy-mm-dd
            return date(int(groups[3]), int(groups[4]), int(groups[5]))
        if groups[6]:  # dd monthname yyyy
            day = int(groups[6])
            month = _FRENCH_MONTHS.get(_normalize_name(groups[7]), 0)
            year = int(groups[8])
            if month:
                return date(year, month, day)
    except (ValueError, TypeError):
        pass
    return None


# ---------------------------------------------------------------------------
# Document parsing
# ---------------------------------------------------------------------------


def _get_paragraphs(path: Path) -> list[str]:
    from docx import Document  # noqa: PLC0415

    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Also include Word page-header paragraphs (date is often written there)
    for section in doc.sections:
        try:
            for p in section.header.paragraphs:
                text = p.text.strip()
                if text:
                    paras.append(text)
        except Exception:  # noqa: BLE001
            pass
    return paras


def _get_tables(path: Path) -> list[list[list[str]]]:
    """Return all tables as list[rows[cells]]."""
    from docx import Document  # noqa: PLC0415

    doc = Document(str(path))
    result = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)
        result.append(rows)
    return result


def _extract_client_block(paragraphs: list[str]) -> tuple[str | None, str | None]:
    """
    Extract client name and address, anchored on the first postal code line.

    Walk backwards from the ZIP CITY line to collect the address block
    (up to 4 lines), then identify the name as the first line that is
    neither a street address nor a salutation.
    """
    postal_re = re.compile(r"\b\d{5}\b")

    def _is_skipped(para: str) -> bool:
        _date_prefix_re = re.compile(
            r"^[A-Za-z\u00c0-\u00ff][A-Za-z\u00c0-\u00ff\s-]+,\s*le\s+",
            re.IGNORECASE,
        )
        return bool(
            _SKIP_PARAGRAPHS_RE.match(para)
            or _date_prefix_re.match(para)
        )

    # Locate the first non-header paragraph that contains a 5-digit postal code.
    postal_idx: int | None = None
    for i, para in enumerate(paragraphs):
        if _is_skipped(para):
            continue
        if re.match(r"^facture\s*n", para, re.IGNORECASE):
            break
        if postal_re.search(para):
            postal_idx = i
            break

    if postal_idx is None:
        # Fallback: return the first non-header, non-date paragraph
        for para in paragraphs:
            if not _is_skipped(para):
                return para, None
        return None, None

    # Walk backwards from the postal line (up to 4 lines) to collect the address block.
    above: list[str] = []
    for i in range(postal_idx - 1, max(postal_idx - 5, -1), -1):
        para = paragraphs[i]
        if _is_skipped(para):
            break
        above.insert(0, para)  # keep top-to-bottom order

    # Identify the client name: first line that is not a street and not a salutation.
    client_name: str | None = None
    street_lines: list[str] = []
    for line in above:
        if client_name is None:
            if not _STREET_START_RE.match(line) and not _SALUTATION_RE.match(line):
                client_name = line
        elif _STREET_START_RE.match(line):
            street_lines.append(line)

    city_line = paragraphs[postal_idx]
    address = "\n".join(street_lines + [city_line]) if street_lines else city_line
    return client_name, address


def _extract_lines_and_total(tables: list[list[list[str]]]) -> tuple[list[InvoiceLine], Decimal]:
    """
    Parse invoice lines from tables.
    Expected format: 2-column table (description | price), last row = Total.
    """
    lines: list[InvoiceLine] = []
    total = Decimal("0.00")

    for table_rows in tables:
        for row in table_rows:
            if len(row) < 2:
                continue
            desc = row[0].strip()
            price_text = row[-1].strip()  # last column is always amount

            price = _parse_price(price_text)
            if price is None:
                continue

            if _TOTAL_RE.match(desc):
                total = price
            elif desc:
                lines.append(InvoiceLine(description=desc, amount=price))

    # If total not found from table, sum lines
    if total == Decimal("0.00") and lines:
        total = sum((ln.amount for ln in lines), Decimal("0.00"))

    return lines, total


def parse_docx(path: Path) -> ParsedInvoice | None:
    """Parse a .docx invoice file and return structured data, or None on failure."""
    m = _FILENAME_RE.match(path.name)
    if not m:
        return None
    number = m.group(1)

    try:
        paragraphs = _get_paragraphs(path)
        tables = _get_tables(path)
    except Exception:
        return None

    # Extract date from paragraphs (body + page headers).
    # Only accept a date whose year matches the filename year to avoid false
    # positives from addresses or birth years in the document.
    filename_year = int(number.split("-")[0])
    invoice_date: date | None = None
    for para in paragraphs:
        d = _parse_date(para)
        if d and d.year == filename_year:
            invoice_date = d
            break

    # Fallback: scan table cells (date sometimes lives in a header row)
    if invoice_date is None:
        for table_rows in tables:
            for row in table_rows:
                for cell in row:
                    d = _parse_date(cell)
                    if d and d.year == filename_year:
                        invoice_date = d
                        break
                if invoice_date:
                    break
            if invoice_date:
                break

    # Final fallback: use the file's last-modified date (filesystem metadata)
    if invoice_date is None:
        invoice_date = date.fromtimestamp(path.stat().st_mtime)

    client_name, client_address = _extract_client_block(paragraphs)
    lines, total = _extract_lines_and_total(tables)

    return ParsedInvoice(
        number=number,
        file=path,
        invoice_date=invoice_date,
        client_name=client_name,
        client_address=client_address,
        lines=lines,
        total_amount=total,
    )


# ---------------------------------------------------------------------------
# Database helpers (sync sqlite3)
# ---------------------------------------------------------------------------


def _db_connect(db_path: Path) -> sqlite3.Connection:
    if not db_path.exists():
        print(f"ERROR: Database not found: {db_path}", file=sys.stderr)
        sys.exit(1)
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _invoice_exists(conn: sqlite3.Connection, number: str) -> bool:
    row = conn.execute("SELECT 1 FROM invoices WHERE number = ?", (number,)).fetchone()
    return row is not None


def _get_existing_pdf_path(conn: sqlite3.Connection, number: str) -> str | None:
    """Return the current pdf_path of an existing invoice, or None if unset."""
    row = conn.execute(
        "SELECT pdf_path FROM invoices WHERE number = ?", (number,)
    ).fetchone()
    if row is None:
        return None
    return row["pdf_path"] or None


def _get_invoice_status(conn: sqlite3.Connection, number: str) -> str | None:
    """Return the status of an existing invoice, or None if not found."""
    row = conn.execute(
        "SELECT status FROM invoices WHERE number = ?", (number,)
    ).fetchone()
    if row is None:
        return None
    return row["status"]


def _is_already_processed(
    conn: sqlite3.Connection, invoice_number: str, pdf_dir: Path
) -> bool:
    """
    Check if an invoice has already been fully processed:
    - Invoice exists in DB
    - PDF file exists on disk (if pdf_path is set)
    """
    row = conn.execute(
        "SELECT pdf_path FROM invoices WHERE number = ?", (invoice_number,)
    ).fetchone()
    if row is None:
        return False
    pdf_path = row["pdf_path"]
    if not pdf_path:
        return False
    # Check if the PDF file actually exists
    pdf_full_path = Path(pdf_path)
    return pdf_full_path.exists()


def _attach_pdf_to_invoice(
    conn: sqlite3.Connection, number: str, pdf_path: str
) -> None:
    """Update pdf_path for an existing invoice."""
    conn.execute(
        "UPDATE invoices SET pdf_path = ?, updated_at = datetime('now') WHERE number = ?",
        (pdf_path, number),
    )


def _find_contact(conn: sqlite3.Connection, client_name: str) -> int | None:
    """Find a contact by name (case-insensitive, both 'NOM PRENOM' and 'PRENOM NOM' order)."""
    norm = _normalize_name(client_name)
    rows = conn.execute(
        "SELECT id, nom, prenom FROM contacts WHERE type = 'client'",
    ).fetchall()
    for row in rows:
        nom_full = _normalize_name(f"{row['nom']} {row['prenom'] or ''}").strip()
        prenom_nom = _normalize_name(f"{row['prenom'] or ''} {row['nom']}").strip()
        if norm in (nom_full, prenom_nom):
            return row["id"]
    return None


def _split_name(full_name: str) -> tuple[str, str | None]:
    """
    Split 'PRENOM NOM' into (nom, prenom).
    Convention: last word = nom, rest = prenom.
    Applies title-case.
    """
    parts = full_name.strip().split()
    if len(parts) == 1:
        return parts[0].title(), None
    nom = parts[-1].title()
    prenom = " ".join(parts[:-1]).title()
    return nom, prenom


def _create_contact(
    conn: sqlite3.Connection,
    client_name: str,
    address: str | None,
) -> int:
    """Insert a new CLIENT contact and return its id."""
    nom, prenom = _split_name(client_name)
    conn.execute(
        """
        INSERT INTO contacts
            (type, nom, prenom, adresse, is_active, blocked, created_at, updated_at)
        VALUES ('client', ?, ?, ?, 1, 0, datetime('now'), datetime('now'))
        """,
        (nom, prenom, address),
    )
    return conn.execute("SELECT last_insert_rowid()").fetchone()[0]


def _mark_invoice_as_paid_if_total_positive(
    amount: Decimal,
) -> bool:
    """Check if an invoice should be marked as paid (amount > 0).
    
    For archived invoices, paid_amount is already set to total_amount at creation,
    so no separate payment record is needed. This function just validates the amount.
    
    Returns True if amount is valid (positive), False otherwise.
    """
    return amount > Decimal("0")


def _cleanup_modern_invoice_pdfs(conn: sqlite3.Connection, pdf_dir: Path, dry_run: bool) -> int:
    """
    Remove PDFs and clear pdf_path for all non-archived invoices (modern invoices).
    Returns count of PDFs removed.
    
    Modern invoices (created in Solde) may have attached PDFs, but these will be
    regenerated on-demand. We clean them to ensure no orphans in data/pdfs/.
    """
    cursor = conn.cursor()
    cursor.execute(
        "SELECT number, pdf_path FROM invoices WHERE status != 'archived' AND pdf_path IS NOT NULL"
    )
    
    removed = 0
    pdf_root = pdf_dir.resolve()
    for row in cursor.fetchall():
        number = row["number"]
        pdf_path = row["pdf_path"]
        
        # Try to delete the PDF file
        pdf_file = Path(pdf_path)
        if not pdf_file.is_absolute():
            pdf_file = (Path.cwd() / pdf_file).resolve()
        else:
            pdf_file = pdf_file.resolve()

        try:
            pdf_file.relative_to(pdf_root)
        except ValueError:
            continue

        if pdf_file.exists():
            try:
                if not dry_run:
                    pdf_file.unlink()
                removed += 1
            except Exception:  # noqa: BLE001
                pass
        
        # Clear pdf_path in DB
        if not dry_run:
            conn.execute(
                "UPDATE invoices SET pdf_path = NULL, updated_at = datetime('now') WHERE number = ?",
                (number,),
            )
    
    if not dry_run:
        conn.commit()
    
    return removed


def _get_next_invoice_seq(conn: sqlite3.Connection, year: int) -> int:
    """Return the next available client invoice sequence number for the given year."""
    prefix = f"{year}-"
    row = conn.execute(
        """
        SELECT number FROM invoices
        WHERE type = 'client' AND number LIKE ?
        ORDER BY number DESC LIMIT 1
        """,
        (f"{prefix}%",),
    ).fetchone()
    if row is None:
        return 1
    last = row["number"]
    try:
        return int(last.split("-")[-1]) + 1
    except (ValueError, IndexError):
        return 1


def _create_invoice(
    conn: sqlite3.Connection,
    parsed: ParsedInvoice,
    contact_id: int,
    pdf_path: str,
) -> int:
    """Insert the invoice and its lines. Returns the new invoice id.
    
    For imported archived invoices, paid_amount is set to total_amount (no payment records needed).
    """
    invoice_date = parsed.invoice_date or date.today()
    number = parsed.number

    conn.execute(
        """
        INSERT INTO invoices (
            number, type, contact_id, date, status, total_amount, paid_amount,
            has_explicit_breakdown, pdf_path, file_path, created_at, updated_at
        ) VALUES (?, 'client', ?, ?, 'archived', ?, ?, ?, ?, NULL, datetime('now'), datetime('now'))
        """,
        (
            number,
            contact_id,
            invoice_date.isoformat(),
            str(parsed.total_amount),
            str(parsed.total_amount),  # Set paid_amount = total_amount for archives
            1 if len(parsed.lines) > 1 else 0,
            pdf_path,
        ),
    )
    invoice_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]

    for ln in parsed.lines:
        conn.execute(
            """
            INSERT INTO invoice_lines
                (invoice_id, description, line_type, quantity, unit_price, amount)
            VALUES (?, ?, 'autres', 1, ?, ?)
            """,
            (invoice_id, ln.description, str(ln.amount), str(ln.amount)),
        )

    return invoice_id


def _find_soffice_binary() -> str | None:
    """Return the LibreOffice CLI binary path if available."""
    return shutil.which("soffice") or shutil.which("libreoffice")


def _convert_docx_to_pdf(
    src: Path,
    invoice_number: str,
    pdf_dir: Path,
    dry_run: bool,
) -> str:
    """Convert a DOCX invoice to PDF and return the stored PDF path (relative to project root).
    
    Skips generation if the PDF file already exists on disk.
    """
    safe_number = invoice_number.replace("/", "-").replace("\\", "-")
    output_path = pdf_dir / f"facture_{safe_number}.pdf"

    # Skip if PDF already exists on disk
    if output_path.exists():
        return str(output_path.relative_to(Path.cwd())).replace("\\", "/")

    # If a sibling PDF already exists, prefer it.
    sibling_pdf = src.with_suffix(".pdf")
    if sibling_pdf.exists():
        if not dry_run:
            pdf_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(sibling_pdf, output_path)
        # Return relative path with forward slashes for cross-platform compatibility
        return str(output_path.relative_to(Path.cwd())).replace("\\", "/")

    # In dry-run mode, do not require LibreOffice; only simulate the target path.
    if dry_run:
        return str(output_path.relative_to(Path.cwd())).replace("\\", "/")

    soffice_bin = _find_soffice_binary()
    if soffice_bin is None:
        raise RuntimeError(
            "LibreOffice CLI not found (soffice). Install LibreOffice or provide sibling .pdf files."
        )

    pdf_dir.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        cmd = [
            soffice_bin,
            "--headless",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(tmp_path),
            str(src),
        ]
        proc = subprocess.run(cmd, check=False, capture_output=True, text=True)
        if proc.returncode != 0:
            stderr = (proc.stderr or "").strip()
            stdout = (proc.stdout or "").strip()
            raise RuntimeError(
                f"DOCX->PDF conversion failed for {src.name}: {stderr or stdout or 'unknown error'}"
            )

        converted_pdf = tmp_path / f"{src.stem}.pdf"
        if not converted_pdf.exists():
            raise RuntimeError(
                f"DOCX->PDF conversion did not produce expected file for {src.name}"
            )
        shutil.copy2(converted_pdf, output_path)

    # Return relative path with forward slashes for cross-platform compatibility
    return str(output_path.relative_to(Path.cwd())).replace("\\", "/")


# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------


def _write_excel_report(report: ImportReport, out_path: Path, dry_run: bool) -> None:
    """Write an Excel report: detail by contact/year, totals by year, new contacts."""
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        print(
            "WARNING: openpyxl non installé — rapport Excel ignoré. "
            "Installez-le avec : pip install openpyxl",
            file=sys.stderr,
        )
        return

    from collections import defaultdict

    records = report.invoice_records
    if not records:
        return

    # ---- Aggregation ----
    detail: dict[tuple[str, int], dict[str, Any]] = defaultdict(
        lambda: {"count": 0, "amount": Decimal("0.00"), "new": False}
    )
    for rec in records:
        key = (rec.client_name, rec.year)
        detail[key]["count"] += 1
        detail[key]["amount"] += rec.amount
        if rec.contact_created:
            detail[key]["new"] = True

    year_agg: dict[int, dict[str, Any]] = defaultdict(
        lambda: {"contacts": set(), "count": 0, "amount": Decimal("0.00")}
    )
    for rec in records:
        year_agg[rec.year]["contacts"].add(rec.client_name)
        year_agg[rec.year]["count"] += 1
        year_agg[rec.year]["amount"] += rec.amount

    # ---- Styles ----
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2E6DA4")
    total_font = Font(bold=True)
    total_fill = PatternFill("solid", fgColor="D9E1F2")
    new_fill = PatternFill("solid", fgColor="E2EFDA")
    center = Alignment(horizontal="center")
    euro_fmt = '#,##0.00 "\u20ac"'

    def _header(ws: Any, headers: list[str], col_widths: list[int]) -> None:
        ws.append(headers)
        row_idx = ws.max_row
        for i, cell in enumerate(ws[row_idx], 1):
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center
            ws.column_dimensions[get_column_letter(i)].width = col_widths[i - 1]

    def _total(ws: Any, values: list[Any]) -> None:
        ws.append(values)
        row_idx = ws.max_row
        for cell in ws[row_idx]:
            cell.font = total_font
            cell.fill = total_fill

    wb = openpyxl.Workbook()

    # ---- Sheet 1: Détail par contact et année ----
    ws1 = wb.active
    ws1.title = "Détail"
    mode_label = "Simulation (dry-run)" if dry_run else "Import réel"
    ws1.append([f"Rapport d'import Word \u2014 {mode_label}"])
    ws1["A1"].font = Font(bold=True, size=12)
    _header(
        ws1,
        ["Contact", "Nouveau", "Année", "Nb factures", "Montant total (\u20ac)"],
        [32, 10, 8, 14, 20],
    )
    sorted_detail = sorted(detail.keys(), key=lambda k: (_normalize_name(k[0]), k[1]))
    for contact, year in sorted_detail:
        data = detail[(contact, year)]
        ws1.append([
            contact,
            "\u2713" if data["new"] else "",
            year,
            data["count"],
            float(data["amount"]),
        ])
        r = ws1.max_row
        ws1.cell(r, 2).alignment = center
        ws1.cell(r, 3).alignment = center
        ws1.cell(r, 4).alignment = center
        ws1.cell(r, 5).number_format = euro_fmt
        if data["new"]:
            for c in range(1, 6):
                ws1.cell(r, c).fill = new_fill
    total_count = sum(d["count"] for d in detail.values())
    total_amount = sum(d["amount"] for d in detail.values())
    _total(ws1, ["TOTAL", "", "", total_count, float(total_amount)])
    ws1.cell(ws1.max_row, 5).number_format = euro_fmt

    # ---- Sheet 2: Par année ----
    ws2 = wb.create_sheet("Par ann\u00e9e")
    _header(
        ws2,
        ["Ann\u00e9e", "Nb contacts", "Nb factures", "Montant total (\u20ac)"],
        [8, 14, 14, 20],
    )
    for year in sorted(year_agg.keys()):
        data = year_agg[year]
        ws2.append([year, len(data["contacts"]), data["count"], float(data["amount"])])
        r = ws2.max_row
        for c in range(1, 4):
            ws2.cell(r, c).alignment = center
        ws2.cell(r, 4).number_format = euro_fmt
    all_contacts = {rec.client_name for rec in records}
    grand_total_amount = sum(d["amount"] for d in year_agg.values())
    grand_total_count = sum(d["count"] for d in year_agg.values())
    _total(ws2, ["TOTAL", len(all_contacts), grand_total_count, float(grand_total_amount)])
    ws2.cell(ws2.max_row, 4).number_format = euro_fmt

    # ---- Sheet 3: Contacts créés ----
    ws3 = wb.create_sheet("Contacts cr\u00e9\u00e9s")
    ws3.append(["Contacts qui seraient cr\u00e9\u00e9s" if dry_run else "Contacts cr\u00e9\u00e9s"])
    ws3["A1"].font = Font(bold=True, size=11)
    _header(ws3, ["Contact", "Nb factures", "Montant total (\u20ac)"], [32, 14, 20])
    new_contact_data: dict[str, dict[str, Any]] = defaultdict(
        lambda: {"count": 0, "amount": Decimal("0.00")}
    )
    for rec in records:
        if rec.contact_created:
            new_contact_data[rec.client_name]["count"] += 1
            new_contact_data[rec.client_name]["amount"] += rec.amount
    if new_contact_data:
        for contact in sorted(new_contact_data, key=_normalize_name):
            data = new_contact_data[contact]
            ws3.append([contact, data["count"], float(data["amount"])])
            r = ws3.max_row
            ws3.cell(r, 2).alignment = center
            ws3.cell(r, 3).number_format = euro_fmt
    else:
        ws3.append(["(aucun nouveau contact)"])

    # ---- Sheet 4: Toutes les factures ----
    ws4 = wb.create_sheet("Factures")
    _header(
        ws4,
        [
            "N\u00b0 facture", "Contact", "Nouveau", "Date",
            "Ann\u00e9e", "Nb lignes", "Montant (\u20ac)", "Fichier source",
        ],
        [14, 32, 10, 14, 8, 10, 18, 40],
    )
    date_fmt = "DD/MM/YYYY"
    for rec in sorted(records, key=lambda r: r.number):
        ws4.append([
            rec.number,
            rec.client_name,
            "\u2713" if rec.contact_created else "",
            rec.invoice_date if rec.invoice_date else "Non trouv\u00e9e",
            rec.year,
            rec.nb_lines,
            float(rec.amount),
            rec.filename,
        ])
        r = ws4.max_row
        ws4.cell(r, 3).alignment = center
        ws4.cell(r, 5).alignment = center
        ws4.cell(r, 6).alignment = center
        ws4.cell(r, 7).number_format = euro_fmt
        if rec.invoice_date:
            ws4.cell(r, 4).number_format = date_fmt
        if rec.contact_created:
            for c in (1, 2, 3):
                ws4.cell(r, c).fill = new_fill
    _total(ws4, ["TOTAL", "", "", "", "", "", float(sum(rec.amount for rec in records)), ""])
    ws4.cell(ws4.max_row, 7).number_format = euro_fmt

    wb.save(str(out_path))
    print(f"\nRapport Excel g\u00e9n\u00e9r\u00e9 : {out_path}")


# ---------------------------------------------------------------------------
# Reconciliation
# ---------------------------------------------------------------------------


def reconcile_orphan_pdfs(db_path: Path, pdf_dir: Path, dry_run: bool) -> None:
    """
    Reconcile orphan PDFs in data/pdfs/ with invoices in DB.
    For each PDF file found, match it to an invoice and update pdf_path if missing.
    Also convert existing absolute paths to relative paths for backend compatibility.
    """
    conn = _db_connect(db_path)
    
    # Get ALL PDF files, not just facture_*.pdf pattern
    all_pdf_files = sorted(pdf_dir.glob("*.pdf"))
    if not all_pdf_files:
        print("No PDF files found in data/pdfs/")
        conn.close()
        return
    
    print(f"Found {len(all_pdf_files)} PDF file(s) in {pdf_dir}")
    
    # Filter files that match the naming pattern
    pdf_files = [f for f in all_pdf_files if re.match(r"^facture_\d{4}-\d{4}\.pdf$", f.name)]
    print(f"  → {len(pdf_files)} match pattern 'facture_YYYY-NNNN.pdf'")
    print(f"  → {len(all_pdf_files) - len(pdf_files)} have other names")
    
    attached = 0
    already_attached = 0
    orphan = 0
    errors = 0
    converted_paths = 0
    
    with Progress(
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("{task.completed}/{task.total}"),
        TimeRemainingColumn(),
    ) as progress:
        task_id = progress.add_task("[cyan]Reconciling PDFs...", total=len(pdf_files))
        
        for pdf_file in pdf_files:
            # Extract invoice number from filename: facture_YYYY-NNNN.pdf
            match = re.match(r"^facture_(\d{4}-\d{4})\.pdf$", pdf_file.name)
            if not match:
                progress.advance(task_id)
                continue
            
            invoice_number = match.group(1)
            pdf_path_abs = str(pdf_file)
            # Use forward slashes for cross-platform compatibility
            pdf_path_rel = str(pdf_file.relative_to(Path.cwd())).replace("\\", "/")
            
            # Check if invoice exists and current pdf_path
            row = conn.execute(
                "SELECT id, pdf_path FROM invoices WHERE number = ?", (invoice_number,)
            ).fetchone()
            
            if row is None:
                orphan += 1
                progress.advance(task_id)
                continue
            
            current_pdf_path = row["pdf_path"]
            
            if current_pdf_path:
                # Invoice already has a pdf_path
                already_attached += 1
                
                # Check if path is absolute and should be converted to relative
                if Path(current_pdf_path).is_absolute() and not dry_run:
                    try:
                        relative_path = str(Path(current_pdf_path).relative_to(Path.cwd()))
                        conn.execute(
                            "UPDATE invoices SET pdf_path = ?, updated_at = datetime('now') WHERE number = ?",
                            (relative_path, invoice_number),
                        )
                        conn.commit()
                        converted_paths += 1
                    except (ValueError, sqlite3.Error):
                        pass  # Path conversion failed, keep as-is
                
                progress.advance(task_id)
                continue
            
            # No pdf_path set: attach the PDF with relative path
            if not dry_run:
                try:
                    conn.execute(
                        "UPDATE invoices SET pdf_path = ?, updated_at = datetime('now') WHERE number = ?",
                        (pdf_path_rel, invoice_number),
                    )
                    conn.commit()
                    attached += 1
                except sqlite3.Error as exc:
                    errors += 1
                    print(f"  ERROR attaching {invoice_number}: {exc}")
            else:
                attached += 1
            
            progress.advance(task_id)
    
    conn.close()
    
    # Summary
    print(f"\n{'=' * 60}")
    print(f"{'DRY-RUN ' if dry_run else ''}RECONCILIATION REPORT")
    print(f"{'=' * 60}")
    verb = "Would attach" if dry_run else "Attached"
    print(f"  {verb:<22}: {attached} PDF(s)")
    print(f"  Already attached        : {already_attached}")
    if converted_paths > 0:
        print(f"  Paths converted (abs→rel): {converted_paths}")
    print(f"  Orphan (no invoice)     : {orphan}")
    print(f"  Errors                  : {errors}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    from rich.progress import Progress, BarColumn, TextColumn, TimeRemainingColumn

    parser = argparse.ArgumentParser(
        description="Import historical Word invoices into Solde (dry-run by default)."
    )
    parser.add_argument("--source", help="Path to folder containing .docx files")
    parser.add_argument("--db", default="data/solde.db", help="Path to SQLite database")
    parser.add_argument("--commit", action="store_true",
                        help="Write changes to DB (default: dry-run)")
    parser.add_argument("--verbose", action="store_true", help="Print detailed info per file")
    parser.add_argument(
        "--report",
        metavar="PATH",
        default="import_report.xlsx",
        help="Output path for the Excel summary report (default: import_report.xlsx)",
    )
    parser.add_argument(
        "--no-report",
        action="store_true",
        dest="no_report",
        help="Skip generating the Excel report",
    )
    parser.add_argument(
        "--reconcile",
        action="store_true",
        help="Reconcile orphan PDFs in data/pdfs/ with invoices in DB (ignores --source)",
    )
    args = parser.parse_args()

    _require_docx()

    # Handle reconciliation mode
    if args.reconcile:
        db_path = Path(args.db)
        pdf_dir = Path("data/pdfs").resolve()
        dry_run = not args.commit
        if dry_run:
            print("DRY-RUN mode — no changes will be made. Pass --commit to write to DB.\n")
        reconcile_orphan_pdfs(db_path, pdf_dir, dry_run)
        return

    # Normal import mode
    if not args.source:
        print("ERROR: --source is required unless --reconcile is used", file=sys.stderr)
        sys.exit(1)

    source = Path(args.source)
    if not source.is_dir():
        print(f"ERROR: Source folder not found: {source}", file=sys.stderr)
        sys.exit(1)

    db_path = Path(args.db)
    pdf_dir = Path("data/pdfs").resolve()
    dry_run = not args.commit

    if dry_run:
        print("DRY-RUN mode — no changes will be made. Pass --commit to write to DB.\n")

    conn = _db_connect(db_path)
    report = ImportReport()

    docx_files = sorted(source.glob("*.docx"))
    print(f"Found {len(docx_files)} .docx file(s) in {source}\n")

    # Cleanup: remove PDFs of modern (non-archived) invoices before import
    print("Cleaning PDFs of modern invoices...", end=" ")
    removed_count = _cleanup_modern_invoice_pdfs(conn, pdf_dir, dry_run)
    print(f"({removed_count} PDF(s) removed)")
    print()

    # Progress bar setup
    progress_columns = [
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("{task.completed}/{task.total}"),
        TimeRemainingColumn(),
    ]

    with Progress(*progress_columns) as progress:
        task_id = progress.add_task("[cyan]Importing invoices...", total=len(docx_files))

        for docx_file in docx_files:
            progress.update(task_id, description=f"[cyan]Processing {docx_file.name}")

            parsed = parse_docx(docx_file)
            if parsed is None:
                if args.verbose:
                    print("  -> Skipped (name does not match pattern)")
                progress.advance(task_id)
                continue

            # Check if file is already fully processed (invoice + PDF both exist on disk)
            if _is_already_processed(conn, parsed.number, pdf_dir):
                report.skipped_already_processed.append(parsed.number)
                if args.verbose:
                    print(f"  -> Skipped: invoice {parsed.number} already processed (invoice + PDF both exist)")
                progress.advance(task_id)
                continue

            # Check if invoice already exists in DB
            if _invoice_exists(conn, parsed.number):
                status = _get_invoice_status(conn, parsed.number)
                
                # If it's a modern invoice (non-archived), skip entirely
                if status != 'archived':
                    report.skipped_exists.append(parsed.number)
                    if args.verbose:
                        print(f"  -> Skipped: invoice {parsed.number} exists as modern ({status}), not importing")
                    progress.advance(task_id)
                    continue
                
                # Archived invoice: check if it has a PDF already
                existing_pdf = _get_existing_pdf_path(conn, parsed.number)
                if existing_pdf:
                    report.skipped_exists.append(parsed.number)
                    if args.verbose:
                        print(f"  -> Skipped: archived invoice {parsed.number} already has PDF")
                    progress.advance(task_id)
                    continue
                
                # Archived invoice without PDF: generate and attach it
                try:
                    pdf_path = _convert_docx_to_pdf(docx_file, parsed.number, pdf_dir, dry_run)
                except Exception as exc:  # noqa: BLE001
                    report.errors.append((parsed.number, str(exc)))
                    if args.verbose:
                        print(f"  -> ERROR converting PDF: {exc}")
                    progress.advance(task_id)
                    continue
                if not dry_run:
                    try:
                        _attach_pdf_to_invoice(conn, parsed.number, pdf_path)
                        conn.commit()
                        report.pdf_attached.append(parsed.number)
                    except sqlite3.Error as exc:
                        conn.rollback()
                        report.errors.append((parsed.number, str(exc)))
                        if args.verbose:
                            print(f"  -> ERROR attaching PDF: {exc}")
                        progress.advance(task_id)
                        continue
                else:
                    report.pdf_attached.append(parsed.number)
                if args.verbose:
                    action = "Would attach" if dry_run else "Attached"
                    print(f"  -> {action} PDF to archived invoice {parsed.number}")
                progress.advance(task_id)
                continue

            # Resolve or create contact
            if parsed.client_name is None:
                report.skipped_no_client.append(parsed.number)
                if args.verbose:
                    print("  -> Skipped: could not extract client name")
                progress.advance(task_id)
                continue

            contact_id = _find_contact(conn, parsed.client_name)
            created_contact = False
            if contact_id is None:
                if dry_run:
                    contact_id = -1
                else:
                    contact_id = _create_contact(conn, parsed.client_name, parsed.client_address)
                created_contact = True

            if args.verbose:
                if not created_contact:
                    contact_info = f"contact #{contact_id}"
                elif dry_run:
                    contact_info = f"would create contact for '{parsed.client_name}'"
                else:
                    contact_info = f"new contact for '{parsed.client_name}'"
                print(
                    f"  -> Date: {parsed.invoice_date}  Client: {parsed.client_name}  "
                    f"Total: {parsed.total_amount}€  Lines: {len(parsed.lines)}  "
                    f"Contact: {contact_info}"
                )

            # Convert DOCX to PDF
            try:
                pdf_path = _convert_docx_to_pdf(docx_file, parsed.number, pdf_dir, dry_run)
            except Exception as exc:  # noqa: BLE001
                report.errors.append((parsed.number, str(exc)))
                if args.verbose:
                    print(f"  -> ERROR converting PDF: {exc}")
                progress.advance(task_id)
                continue

            # Create invoice
            if not dry_run:
                try:
                    # Invoice amount must be strictly positive
                    if not _mark_invoice_as_paid_if_total_positive(parsed.total_amount):
                        report.errors.append((parsed.number, "Invoice amount must be > 0"))
                        progress.advance(task_id)
                        continue
                    
                    invoice_id = _create_invoice(conn, parsed, contact_id, pdf_path)
                    invoice_date = parsed.invoice_date or date.today()
                    # For archived invoices, paid_amount is already set at creation.
                    # No separate payment record is created.
                    conn.commit()
                    report.created.append(parsed.number)
                    report.invoice_records.append(InvoiceRecord(
                        number=parsed.number,
                        client_name=parsed.client_name or "",
                        contact_created=created_contact,
                        year=invoice_date.year,
                        amount=parsed.total_amount,
                        invoice_date=parsed.invoice_date,
                        filename=docx_file.name,
                        nb_lines=len(parsed.lines),
                    ))
                except sqlite3.IntegrityError as exc:
                    conn.rollback()
                    report.errors.append((parsed.number, str(exc)))
                    if args.verbose:
                        print(f"  -> ERROR: {exc}")
            else:
                report.created.append(parsed.number)
                report.invoice_records.append(InvoiceRecord(
                    number=parsed.number,
                    client_name=parsed.client_name or "",
                    contact_created=created_contact,
                    year=(parsed.invoice_date or date.today()).year,
                    amount=parsed.total_amount,
                    invoice_date=parsed.invoice_date,
                    filename=docx_file.name,
                    nb_lines=len(parsed.lines),
                ))

            progress.advance(task_id)

    conn.close()

    # Summary
    print(f"\n{'=' * 60}")
    print(f"{'DRY-RUN ' if dry_run else ''}IMPORT REPORT")
    print(f"{'=' * 60}")
    verb_create = "Would create" if dry_run else "Created"
    verb_attach = "Would attach" if dry_run else "Attached"
    print(f"  {verb_create:<22}: {len(report.created)} invoice(s)")
    print(f"  {verb_attach} PDF        : {len(report.pdf_attached)} invoice(s)")
    print(f"  Skipped (already processed): {len(report.skipped_already_processed)}")
    print(f"  Skipped (already exists)  : {len(report.skipped_exists)}")
    print(f"  Skipped (no client name)  : {len(report.skipped_no_client)}")
    print(f"  Errors                    : {len(report.errors)}")
    if report.skipped_no_client:
        print("\nSkipped (no client name):")
        for number in report.skipped_no_client:
            print(f"  - {number}: client name could not be extracted")
    if report.errors:
        print("\nErrors:")
        for number, msg in report.errors:
            print(f"  - {number}: {msg}")
    if dry_run and (report.created or report.pdf_attached):
        total_dry = len(report.created) + len(report.pdf_attached)
        print(f"\nRun with --commit to actually import/attach PDF for {total_dry} invoice(s).")

    if not args.no_report:
        _write_excel_report(report, Path(args.report), dry_run)


if __name__ == "__main__":
    main()
